import subprocess
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement

# Función para convertir SVG a PNG usando Inkscape
def convert_svg_to_png(svg_path, png_path):
    subprocess.run(['inkscape', svg_path, '--export-type=png', '-o', png_path])

# Lista de íconos y sus archivos SVG correspondientes
icon_files = {
    'phone': 'telephone.svg',
    'address': 'geo-alt.svg',
    'email': 'envelope-at.svg',
    'linkedin': 'linkedin.svg',
    'github': 'github.svg',
    'portfolio': 'file-person.svg',
    'education': 'education.svg',
    'experience': 'work.svg',
    'skills': 'skills.svg',
    'certification': 'certification.svg'
}

# Convertir todos los archivos SVG a PNG
for key, svg_file in icon_files.items():
    png_file = f"{key}.png"
    convert_svg_to_png(svg_file, png_file)

# Crear un nuevo documento
doc = Document()

# Establecer fuente predeterminada
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Configurar márgenes del documento
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Añadir foto
doc.add_picture('image.png', width=Inches(1.5))
last_paragraph = doc.paragraphs[-1] 
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Título
title = doc.add_heading('Luis Carlos Gallardo Ramirez', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.runs[0]
run.font.size = Pt(28)
run.font.name = 'Calibri'
run.bold = True

# Subtítulo
subtitle = doc.add_paragraph('IT System Administrator')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.italic = True

# Añadir espacio después del título y subtítulo
doc.add_paragraph("\n")

# Información de contacto con logos
table = doc.add_table(rows=2, cols=2)
table.autofit = True
table.columns[0].width = Inches(3)
table.columns[1].width = Inches(3)
hdr_cells = table.rows[0].cells

# Añadir íconos y detalles de contacto en una celda
contact_info = hdr_cells[0].paragraphs[0]
contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

contacts = [
    ('phone.png', '+52 (668) 136-9109'),
    ('geo-alt.png', 'Los Mochis, Sinaloa, Mexico'),
    ('envelope-at.png', 'Luissgr95@gmail.com')
]

for icon, info in contacts:
    contact_info.add_run('\n')  # Añadir salto de línea antes del ícono
    run = contact_info.add_run()
    run.add_picture(icon, width=Inches(0.15))
    run.add_text(f' {info}')

# Redes sociales con íconos en otra celda
social_media = hdr_cells[1].paragraphs[0]
social_media.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

logos = [
    ('linkedin.png', 'www.linkedin.com/in/luisgr95/'),
    ('github.png', 'www.github.com/luisgr95'),
    ('portfolio.png', 'www.luisgr95.ga')
]

for logo, url in logos:
    social_media.add_run('\n')  # Añadir salto de línea antes del ícono
    run = social_media.add_run()
    run.add_picture(logo, width=Inches(0.15))
    run.add_text(f' {url}')

# Asegurar que los íconos y textos de redes sociales estén alineados
for paragraph in hdr_cells[1].paragraphs:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)

# Añadir línea separadora
def add_separator():
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break()
    run.font.size = Pt(1)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run.add_text("_" * 100)
    run.add_break()
    run.add_text("\n")

# Añadir sección de experiencia laboral
def add_section_heading(heading, icon_path):
    heading_paragraph = doc.add_paragraph()
    heading_paragraph.paragraph_format.space_before = Pt(12)
    heading_run = heading_paragraph.add_run()
    heading_run.add_picture(icon_path, width=Inches(0.2))
    heading_run.add_text(f' {heading}')
    heading_run.font.size = Pt(16)
    heading_run.bold = True
    add_separator()

add_section_heading('Work Experience', 'experience.png')

# Experiencia en IBM
doc.add_heading('Technical Support Professional - System Administrator', level=2)
exp_1 = doc.add_paragraph('IBM\nNovember 2022 - Present\nGuadalajara, Jalisco', style='Intense Quote')
exp_1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

ibm_experience = doc.add_paragraph()
ibm_experience.paragraph_format.line_spacing = Pt(12)
ibm_experience.add_run(
    """
    - Provide technical support for multiple systems, including z/OS,z/VM and cloud environments.
    - Experience in ticketing systems, troubleshooting support tickets in assigned queues(ServiceNow,AccessHub,PagerDuty).
    - Development and management of infrastructure using Microsoft Azure & IBM Cloud for various projects.
    - Assist with cloud integration and automation tasks.
    - Automate processes and improve support workflows using Python and Ansible.
    - SQL proficiency for database management and report generation(Query Scripting).
    - Work with APIs such as IBM Capki, W3 People, Ansible, GitHub, and ServiceNow for seamless cloud integrations.
    - Create and maintain digital certificates.
    - Work in general security issues and provide support to other areas like CICS, DB2, IMS, etc.
    - Implement new procedures or processes independently or through the team to solve problems, analyze situations and implement solutions as needed to remediate issues.
    - Participate in audit initiatives and provide evidence, supporting tasks and work in non-compliance issues.
    - Support system health check reviewing TechSpec and providing documentation.
    - Development of scripts in Python to manipulate reports generated by zSecure, formatting them in CSV or manipulating them as needed, and executing them from Ansible.
    - Implementation of DevOps methodologies using CI/CD with GitHub Actions.
    - Use of Grafana with Prometheus to monitor system resources.
    - Development of Ansible scripts to execute tasks on z/OS servers related to logical security, audits, reports, certificate management, and user reports.
    - Creation of visual tools with Tkinter to facilitate tasks for other teams within the same management.
    """
)

# Añadir espacio después de la sección
doc.add_paragraph()

# Experiencia en Universidad del Valle del Fuerte
doc.add_heading('IT Manager', level=2)
exp_2 = doc.add_paragraph('Universidad del Valle del Fuerte S.C\nSeptember 2021 - November 2022\nLos Mochis, Sinaloa, Mexico', style='Intense Quote')
exp_2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

univ_experience = doc.add_paragraph()
univ_experience.paragraph_format.line_spacing = Pt(12)
univ_experience.add_run(
    """
    - Managed IT infrastructure, including cloud solutions and on-premises servers.
    - Developed automation scripts for infrastructure management, reducing operational load.
    - Worked with Microsoft Azure to deploy and manage cloud resources.
    - Implemented ticketing and support systems to improve response times and efficiency.
    - Set up and maintained networks, servers, and systems.
    - Deployment in production Server Windows Server also configuring Active Directory, assigning GPO policies, shared folder service and LAPS.
    - Deployment in production Linux CentOS Server also configuring Apache, PHP, Mysql for hosting the Website and ERP Odoo.
    - Develop addons for ERP Odoo with Python and PostgreSQL.
    - Installation, Configuration, Administration and Security of the Network Infrastructure using technologies such as Cisco, MikroTik, Ubiquiti, GrandStream and Fortinet.
    - Development and Implementation of the RED UNIVAFU website using angular as framework and Django for the backend implementing mysql for the database.
    - Communicated with clients and workers regarding appointment scheduling and technical support.
    - Recommended changes, additions, and removals of software to increase operating efficiency by up to 47%.
    - Developed expertise through 12+ hours of training each month in cutting-edge technology, including networking, servers, ERP management, and firewalls.
    - Professor of the Engineering Career in Computer Systems and Digital Marketing teaching the subjects of networks and services GNU/Linux.
    """
)

# Experiencia en Comercializadora Gisnet S.A de CV
doc.add_heading('IT Technician', level=2)
exp_3 = doc.add_paragraph('Comercializadora Gisnet S.A de CV\nMay 2021 - September 2021\nSinaloa, Mexico', style='Intense Quote')
exp_3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

gisnet_experience = doc.add_paragraph()
gisnet_experience.paragraph_format.line_spacing = Pt(12)
gisnet_experience.add_run(
    """
    - Installed software and performed system updates for Santander y BBVA systems.
    - Collaborated with IT team to respond to work tickets using Jira, resolving urgent issues, documenting ticket resolutions, and completing projects on time.
    - Communicated with clients regarding appointment scheduling and technical support and traveled within a 250-km radius to provide on-site work.
    """
)

# Experiencia en Nueva Senda College
doc.add_heading('IT Technician and Technology Teacher', level=2)
exp_4 = doc.add_paragraph('Nueva Senda College\nSeptember 2020 - May 2021\nLos Mochis, Sinaloa, Mexico', style='Intense Quote')
exp_4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

nueva_senda_experience = doc.add_paragraph()
nueva_senda_experience.paragraph_format.line_spacing = Pt(12)
nueva_senda_experience.add_run(
    """
    - Organization, planning and coordination of subjects.
    - Provide theoretical-practical teaching to students.
    - Digital platform management.
    - Networks, CCTV-IP, Servers, Preventive and corrective maintenance.
    """
)

# Experiencia en AGROEQUIPOS DEL VALLE S.A DE C.V
doc.add_heading('Professional Practices', level=2)
exp_5 = doc.add_paragraph('AGROEQUIPOS DEL VALLE S.A DE C.V\nJanuary 2020 - July 2020\nLos Mochis, Sinaloa, Mexico', style='Intense Quote')
exp_5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

agroequipos_experience = doc.add_paragraph()
agroequipos_experience.paragraph_format.line_spacing = Pt(12)
agroequipos_experience.add_run(
    """
    - Development of "Data Analysis" project with Power BI and VBA.
    - Installation of GPS equipment on tractors to create a precision farming environment.
    """
)

add_separator()

# Añadir sección de resumen de carrera
add_section_heading('Career Summary', 'file-person.png')

career_summary = doc.add_paragraph()
career_summary.paragraph_format.line_spacing = Pt(12)
career_summary.add_run(
    """
    I am a systems engineer with over 5 years of experience in the field of Information Technology, including 4 years as a specialized Sysadmin in Cybersecurity. Currently, I lead automation projects using Ansible and Python, significantly optimizing our processes in z/OS and Linux environments. Additionally, I am responsible for key and certificate management, working closely with experts to ensure security and authorized access in my work environment. Recently, I have expanded my experience in programming scripts in Python and Ansible, working with APIs and DevOps methodologies to improve CI/CD processes and system monitoring.
    """
)

add_separator()

# Añadir sección de educación
add_section_heading('Education', 'education.png')

education = doc.add_paragraph()
education.paragraph_format.line_spacing = Pt(12)
education.add_run(
    """
    Bachelor's Degree in Engineering
    Mechatronics Engineering
    Instituto Tecnológico de los Mochis
    2017 - 2021
    Los Mochis, Sinaloa, Mexico

    Master of Engineering
    Software Engineering & Computer Systems
    Universidad de la Rioja
    January 2022 - August 2023
    La Rioja, Spain
    """
)

add_separator()

# Añadir sección de habilidades
add_section_heading('Skills', 'skills.png')

skills = doc.add_paragraph()
skills.paragraph_format.line_spacing = Pt(12)
skills.add_run(
    """
    - Consistent
    - Project Management
    - Problem-solving
    - System Admin (Windows Server, Redhat, z/OS)
    - Strong Communicator
    - Jira
    - Databases (SQL, MySQL, MariaDB)
    - Networking
    - Software Developer
    - SSH
    - Programming Languages (C, Python, JavaScript, JCL, REXX)
    - CMS (Drupal7, Drupal8, Wordpress, Prestashop y Magento)
    - Task Automation
    - Information Security
    - APIs (IBM Capki, W3 People, Ansible, GitHub, ServiceNow)
    - CI/CD (GitHub Actions)
    - Monitoring (Grafana, Prometheus)
    - Visual Tools (Tkinter)
    """
)

add_separator()

# Añadir sección de certificaciones
add_section_heading('Certifications', 'certification.png')

certifications = doc.add_paragraph()
certifications.paragraph_format.line_spacing = Pt(12)
certifications.add_run(
    """
    - Continuing Education Certificate in Digital Project Management
      MIU City University Miami
      April 2024 - April 2030
      Credential ID: 55d47092-f812-4b28-82d8-7a1d40a69443

    - Interskill - Mainframe Systems Auditor – RACF – Experienced 2.4
      IBM
      February 2024

    - Trustworthy AI and AI Ethics
      IBM
      February 2024

    - Interskill - Mainframe Systems Operator – TSO/ISPF Intermediate 2.4
      IBM
      January 2024

    - Think Like a Hacker
      IBM
      January 2024
      Credential ID: 1aba9425-6555-4736-b5f6-c4ab093bfde6

    - Diploma en Seguridad Informática
      UNIR México
      August 2023
      Credential ID: d1d456c8-57db-4e17-a2b8-3ebac459c9bb

    - Generative AI and watsonx for T&O
      IBM
      July 2023
      Credential ID: PLAN-2E22FC00BF9E

    - IBM Agile Explorer
      IBM
      November 2022

    - Interskill - Mainframe Operator – TSO/ISPF Foundations 2.4
      IBM
      November 2022

    - Proyectos de Software
      UNIR México
      October 2022
      Credential ID: 43a8fb66-a116-40ad-b18d-4eadf2b86df0

    - Ingeniería Web
      UNIR México
      August 2022
      Credential ID: da40be37-fd75-49ba-a211-010936d62942

    - Diplomado en Soft Skills y Habilidades Directivas
      Universidad Internacional de la Rioja - España
      June 2022
      Credential ID: 521c5782-775c-4ad4-8d9a-9d5eb14c4e19

    - Certificación en Linux
      Mastermind
      February 2022
      Credential ID: e53202cf132f4b0fbbdaf91cf8de9607
    """
)

# Guardar el documento
doc.save('Updated_CV_LuisGallardo.docx')
